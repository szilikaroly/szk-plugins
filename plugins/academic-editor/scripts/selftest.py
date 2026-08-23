#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Self test for the checker and the two .docx writers. Offline. Exit 0 = pass.

The tracked-changes writer is the part that has to be right: an edit that
silently fails to apply is invisible in the delivered file, and an edit that
mangles a run destroys the author's formatting. So the round-trip is asserted
both ways — reject must reproduce the input text exactly, accept must reproduce
the intended output — on a paragraph whose target phrase is deliberately split
across three differently formatted runs.
"""
from __future__ import annotations

import io
import json
import sys
import tempfile
from contextlib import redirect_stdout
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
import manuscript_check as MC        # noqa: E402


def _ok(label: str, cond: bool) -> bool:
    print(f"  {'PASS' if cond else 'FAIL'}  {label}")
    return bool(cond)


BAD = """Effects of Sertraline on Sleep

Abstract
We assessed 31 patients (n=31) using the Hamilton Rating Scale for Depression (HRSD).
The HRSD scores fell. Firstly, the doses were 200mg/day and so on. Results are in
table 2 and figure 1.

Introduction
The centre reported a statistical difference on the 14thday of treatment.
Analysis used the Independent Ethics Committee (IEC) framework.

Methods
Values are 12.0% ± 4.3% and 5.1±2.3%. Significance was p=0.004 and P=0.005.
A two-sided 5% level of significance was considered statistically significant.
Written informed consents were signed prior to participation.
Behaviour was analysed at the center.

Results
The ESS scores fell. The ESS remained stable. The ESS was low.

Discussion
This matters.

References
1. Someone. A paper. J Sleep. 2020.
"""

CLEAN = """A Short Report

Abstract
We assessed thirty-one patients using a validated scale.

Introduction
The center reported a significant difference on the 14th day of treatment.

Methods
Values are 12.0% ± 4.3%. Significance was set at p=0.05.
Written informed consent was obtained from each patient prior to participation.

Results
Scores fell.

Discussion
This matters.
"""


def _findings(text: str, only: list[str] | None = None) -> list[dict]:
    p = Path(tempfile.mkdtemp()) / "ms.txt"
    p.write_text(text, encoding="utf-8")
    body = MC.read_text(p)
    secs = MC.split_sections(body)
    checks = only or list(MC.CHECKS)
    out = []
    for name in checks:
        out.extend(MC.CHECKS[name](body, secs))
    return out


def _msgs(findings, check: str) -> str:
    return "\n".join(f["msg"] for f in findings if f["check"] == check)


def _tracked_text(path: Path, mode: str) -> str:
    """Read a .docx as the author sees it. mode='orig' rejects insertions."""
    import zipfile
    from lxml import etree
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    doc = etree.fromstring(zipfile.ZipFile(path).read("word/document.xml"))
    out = []
    for p in doc.iter(W + "p"):
        buf = []
        for node in p.iter():
            if node.tag == W + "t":
                if mode == "orig" and any(a.tag == W + "ins"
                                          for a in node.iterancestors()):
                    continue
                buf.append(node.text or "")
            elif node.tag == W + "delText" and mode == "orig":
                buf.append(node.text or "")
        out.append("".join(buf))
    return "\n".join(out)


def run() -> bool:
    ok = True

    print("\n[manuscript_check — sections]")
    f = _findings(BAD, ["counts"])
    ok &= _ok("IMRaD sections detected",
              "abstract" in _msgs(f, "counts") and "methods" in _msgs(f, "counts"))

    print("\n[manuscript_check — abbreviations]")
    f = _findings(BAD, ["abbreviations"])
    m = _msgs(f, "abbreviations")
    ok &= _ok("'IEC' flagged as defined but never used", "'IEC'" in m)
    ok &= _ok("'HRSD' NOT flagged when it is actually used again",
              "'HRSD' definiálva" not in m)
    # A glossary line excuses earlier use — without this the checker produced a
    # 20-line false alarm on any manuscript carrying an abbreviations list.
    with_glossary = BAD.replace("Abstract\n",
                                "Abstract\nAbbreviations: PSG: polysomnography; "
                                "TST: total sleep time\n")
    g = _msgs(_findings(with_glossary, ["abbreviations"]), "abbreviations")
    ok &= _ok("a glossary entry licenses earlier use",
              "'PSG' a definíciója ELŐTT" not in g)

    print("\n[manuscript_check — consistency, reported as splits]")
    f = _findings(BAD, ["consistency"])
    m = _msgs(f, "consistency")
    ok &= _ok("p/P split reported", "`p` és `P`" in m)
    ok &= _ok("± spacing split reported", "`±`" in m)
    ok &= _ok("the p/P finding warns about silent global rewriting",
              "eredmény átírása" in m)
    u = _msgs(f, "units")
    ok &= _ok("missing unit space caught (200mg/day)", "200mg" in u)
    ok &= _ok("run-together ordinal caught, showing the whole token (14thday)",
              "14thday" in u)
    ok &= _ok("lowercase 'table 2' / 'figure 1' caught",
              "table" in _msgs(f, "references-to-floats").lower())

    print("\n[manuscript_check — spelling]")
    f = _findings(BAD, ["spelling"])
    ok &= _ok("UK/US mix caught (centre + center)",
              "UK és US helyesírás keverve" in _msgs(f, "spelling"))
    # "analysis" must not read as the British "analyse".
    neutral = _findings("Methods\nThe analysis of the data was done at the center.\n",
                        ["spelling"])
    ok &= _ok("'analysis' is not counted as UK spelling",
              "keverve" not in _msgs(neutral, "spelling"))

    print("\n[manuscript_check — register markers from the reference edit]")
    f = _findings(BAD, ["register"])
    m = _msgs(f, "register")
    for marker, label in [("and so on", "list-closer"), ("Firstly", "Firstly"),
                          ("statistical difference", "statistical difference"),
                          ("performed by using", None),
                          ("considered statistically significant",
                           "significance LEVEL"),
                          ("consents were signed", "informed consent")]:
        if label:
            ok &= _ok(f"'{marker}' flagged", marker.lower() in m.lower())
    ok &= _ok("a clean manuscript raises no register markers",
              _msgs(_findings(CLEAN, ["register"]), "register") == "")

    print("\n[docx writers]")
    try:
        import docx  # noqa: F401
        from lxml import etree  # noqa: F401
    except ImportError:
        print("  SKIP  python-docx / lxml not installed on this interpreter")
        print(f"\n{'ALL PASSED' if ok else 'FAILURES PRESENT'}")
        return ok

    import docx as _docx
    import docx_tracked_edit as TE
    import docx_accept_changes as AC

    tmp = Path(tempfile.mkdtemp())
    src = tmp / "ms.docx"
    d = _docx.Document()
    d.add_heading("Methods", level=1)
    p = d.add_paragraph("All statistical procedures were performed by using SPSS. ")
    # The target phrase is split across three runs with different formatting —
    # this is what Word actually produces, and what naive replacement breaks.
    p.add_run("A two-sided 5% level ")
    p.add_run("of significance ")
    r = p.add_run("was considered statistically significant.")
    r.bold = True
    d.add_paragraph("The ESS fell. The ESS rose. The ESS fell again.")
    d.save(str(src))

    original = _tracked_text(src, "final")

    edits = [
        {"find": "performed by using SPSS",
         "replace": "performed using SPSS"},
        {"find": "A two-sided 5% level of significance was considered statistically significant.",
         "replace": "A two-sided 5% level of significance was applied.",
         "comment": "A significance level cannot itself be significant; the "
                    "threshold, sidedness and value are unchanged."},
        {"find": "The ESS", "replace": "The ESS score", "count": 0},
        {"find": "THIS PHRASE IS NOT PRESENT", "replace": "x"},
    ]
    ef = tmp / "edits.json"
    ef.write_text(json.dumps(edits, ensure_ascii=False), encoding="utf-8")

    buf = io.StringIO()
    with redirect_stdout(buf):
        rc = TE.main([str(src), str(ef), "--dry-run"])
    dry = buf.getvalue()
    ok &= _ok("--dry-run names the edit that did not match",
              "THIS PHRASE IS NOT PRESENT" in dry)
    ok &= _ok("--dry-run exits non-zero when something missed", rc == 1)

    out = tmp / "ms_edited.docx"
    with redirect_stdout(io.StringIO()):
        TE.main([str(src), str(ef), "-o", str(out)])
    ok &= _ok("tracked file written", out.exists())

    final = _tracked_text(out, "final")
    orig_back = _tracked_text(out, "orig")
    ok &= _ok("a phrase spanning three formatted runs is replaced",
              "was applied." in final and "considered statistically significant"
              not in final)
    ok &= _ok("count:0 replaces every occurrence (3x 'The ESS score')",
              final.count("The ESS score") == 3)
    ok &= _ok("rejecting the insertions reproduces the original text",
              orig_back.strip() == original.strip())

    import zipfile
    z = zipfile.ZipFile(out)
    ok &= _ok("comments part created", "word/comments.xml" in z.namelist())
    ok &= _ok("content-types override added",
              "/word/comments.xml" in z.read("[Content_Types].xml").decode())
    ok &= _ok("relationship added",
              "comments.xml" in z.read("word/_rels/document.xml.rels").decode())
    ok &= _ok("every part is well-formed XML",
              all(_wellformed(z, n) for n in z.namelist()
                  if n.endswith((".xml", ".rels"))))

    clean = tmp / "ms_clean.docx"
    rejected = tmp / "ms_rejected.docx"
    with redirect_stdout(io.StringIO()):
        AC.main([str(out), "-o", str(clean)])
        AC.main([str(out), "--reject", "-o", str(rejected)])
    ok &= _ok("accept reproduces the edited text",
              _tracked_text(clean, "final").strip() == final.strip())
    ok &= _ok("reject reproduces the author's original",
              _tracked_text(rejected, "final").strip() == original.strip())
    ok &= _ok("the clean copy carries no comments part",
              "word/comments.xml" not in zipfile.ZipFile(clean).namelist())
    ok &= _ok("the clean copy still opens with python-docx",
              len(_docx.Document(str(clean)).paragraphs) >= 3)

    print(f"\n{'ALL PASSED' if ok else 'FAILURES PRESENT'}")
    return ok


def _wellformed(z, name: str) -> bool:
    from lxml import etree
    try:
        etree.fromstring(z.read(name))
        return True
    except Exception:
        return False


if __name__ == "__main__":
    sys.exit(0 if run() else 1)
