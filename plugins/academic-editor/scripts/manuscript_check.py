#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Deterministic pre-edit pass. Everything judgement is bad at, and nothing else.

    manuscript_check.py manuscript.docx
    manuscript_check.py manuscript.docx --json report.json
    manuscript_check.py manuscript.docx --only abbreviations,spelling

What it reports, and why each one is here rather than left to reading
--------------------------------------------------------------------
* word counts, whole and by IMRaD section — the number the author needs after
  an edit, and the one thing an editor must never estimate;
* abbreviations defined-but-never-used and used-before-defined — a class of
  defect that is invisible on a linear read and trivial to count;
* sentence-length outliers — not a style verdict, a list of places to look;
* passive density by section — Methods is *supposed* to be passive; Discussion
  usually is not. The number is context, not a target;
* `p` vs `P`, `±` spacing, operator spacing — the reference edit was internally
  INCONSISTENT on all three, and the operative rule is local consistency plus a
  query. So this reports the split and refuses to pick a side;
* US/UK spelling splits — one manuscript, one convention;
* a small set of register markers the reference edit removed on sight.

No language model, no spell dictionary: it never false-positives on a medical
term it has not heard of. Stdlib only, except python-docx for .docx input.
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from collections import Counter
from pathlib import Path

# --------------------------------------------------------------------------- input

SECTION_RE = re.compile(
    r"^\s*(?:\d+(?:\.\d+)*\.?\s*)?"
    r"(abstract|introduction|background|methods?|materials and methods|"
    r"patients and methods|results|discussion|conclusions?|limitations|"
    r"references?|acknowledge?ments?|funding|conflicts? of interest)\b",
    re.I)


def read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ".tex"):
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        try:
            import docx
        except ImportError:
            sys.exit("A .docx olvasásához python-docx kell: pip install python-docx")
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    for tool in ("doctotext", "pdftotext"):
        try:
            out = subprocess.run([tool, str(path)], capture_output=True, text=True,
                                 timeout=120)
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout
        except (OSError, subprocess.SubprocessError):
            continue
    sys.exit(f"Nem tudom kiolvasni: {path} (próbáld a doc-tools skill doctotext-jét)")


def split_sections(text: str) -> dict[str, str]:
    """IMRaD split on heading-looking lines. Text before the first heading is 'front'."""
    sections: dict[str, list[str]] = {"front": []}
    current = "front"
    for line in text.splitlines():
        m = SECTION_RE.match(line.strip())
        if m and len(line.strip()) < 80:
            current = m.group(1).lower()
            current = {"background": "introduction",
                       "materials and methods": "methods",
                       "patients and methods": "methods",
                       "method": "methods",
                       "conclusion": "conclusions",
                       "reference": "references"}.get(current, current)
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


# --------------------------------------------------------------------------- helpers

_ABBR_DEF = re.compile(r"\(([A-Z][A-Za-z0-9\-]{1,9})\)|\[([A-Z][A-Za-z0-9\-]{1,9})\]")
_ABBR_USE = re.compile(r"\b([A-Z]{2,9}(?:-[A-Z0-9]{1,4})?)\b")
_SENT_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z“\"(])")

#: Deliberately short. Every pair is one where a mixed manuscript is a defect,
#: and none of them collides with a drug name, gene symbol or anatomical term.
UK_US = [
    ("centre", "center"), ("colour", "color"), ("behaviour", "behavior"),
    ("analyse", "analyze"), ("analysed", "analyzed"), ("analysing", "analyzing"),
    ("randomised", "randomized"), ("randomisation", "randomization"),
    ("organisation", "organization"), ("recognised", "recognized"),
    ("hospitalised", "hospitalized"), ("summarise", "summarize"),
    ("characterised", "characterized"), ("utilise", "utilize"),
    ("labelled", "labeled"), ("modelling", "modeling"), ("programme", "program"),
    ("oedema", "edema"), ("anaemia", "anemia"), ("haemorrhage", "hemorrhage"),
    ("paediatric", "pediatric"), ("foetal", "fetal"), ("oesophag", "esophag"),
    ("tumour", "tumor"), ("fibre", "fiber"), ("litre", "liter"), ("metre", "meter"),
]

#: Removed on sight in the reference edit. Reported, never auto-changed.
REGISTER_MARKERS = [
    (r"\band so on\b", "list-closer — the reference edit deleted it twice"),
    (r"\bete\.|\betc\.", "`etc.` in a scientific list; enumerate or cut"),
    (r"\bFirstly\b|\bSecondly\b|\bThirdly\b", "→ first / second / third"),
    (r"\bstatistical difference", "→ significant difference"),
    (r"\bthe latter (?:\d+|three|four|several)\b", "`latter` = the second of two; use `last`"),
    (r"\bin this direction\b", "signals no relation; the reference edit deleted it"),
    (r"\bperformed by using\b", "→ performed using"),
    (r"\bwas considered statistically significant\b",
     "a significance LEVEL cannot be significant; → was applied"),
    (r"\bconsents were signed\b", "→ informed consent was obtained from each patient"),
]


def sentences(text: str) -> list[str]:
    body = re.sub(r"\s+", " ", text)
    return [s.strip() for s in _SENT_SPLIT.split(body) if len(s.strip()) > 1]


def words(text: str) -> list[str]:
    return re.findall(r"[A-Za-z0-9][A-Za-z0-9'\-]*", text)


_PASSIVE = re.compile(
    r"\b(?:is|are|was|were|been|being|be)\s+(?:\w+ly\s+)?(\w+(?:ed|en|wn|ne|ung|ought))\b",
    re.I)


# --------------------------------------------------------------------------- checks

def check_counts(sections: dict[str, str]) -> list[dict]:
    out = []
    total = len(words("\n".join(v for k, v in sections.items() if k != "references")))
    out.append({"level": "info", "check": "counts",
                "msg": f"{total} szó (a References nélkül)"})
    for name in ("abstract", "introduction", "methods", "results", "discussion",
                 "conclusions"):
        if sections.get(name):
            n = len(words(sections[name]))
            out.append({"level": "info", "check": "counts",
                        "msg": f"  {name}: {n} szó"})
            if name == "abstract" and n > 300:
                out.append({"level": "warn", "check": "counts",
                            "msg": f"az absztrakt {n} szó — a legtöbb folyóirat 250-300-nál "
                                   "húzza meg; ellenőrizd a célújság limitjét"})
    return out


#: `AASM-2007: American Academy of Sleep Medicine` inside an abbreviations block.
#: A glossary is a definition too — treating only parenthetical expansions as
#: definitions made the checker call every abbreviation "used before defined" in
#: any manuscript carrying the list journals increasingly require, which buries
#: the real findings under twenty false ones.
_ABBR_LIST = re.compile(r"(?:^|[;,]\s*)([A-Z][A-Za-z0-9\-]{1,9})\s*:\s+[A-Za-z]", re.M)


def check_abbreviations(text: str) -> list[dict]:
    inline: dict[str, int] = {}          # expanded in parentheses at some point
    glossary: set[str] = set()           # listed in an abbreviations block
    for m in _ABBR_DEF.finditer(text):
        abbr = m.group(1) or m.group(2)
        if abbr and abbr.upper() == abbr and len(abbr) >= 2:
            inline.setdefault(abbr, m.start())
    for m in _ABBR_LIST.finditer(text):
        abbr = m.group(1)
        if abbr.upper() == abbr and len(abbr) >= 2:
            glossary.add(abbr)

    used = Counter()
    first_use: dict[str, int] = {}
    for m in _ABBR_USE.finditer(text):
        used[m.group(1)] += 1
        first_use.setdefault(m.group(1), m.start())

    out = []
    for abbr in sorted(set(inline) | glossary):
        # The definition itself counts as one occurrence of the token.
        if used[abbr] <= 1:
            out.append({"level": "warn", "check": "abbreviations",
                        "msg": f"'{abbr}' definiálva, de utána nem használva — "
                               "vagy töröld a definíciót, vagy használd"})
    for abbr, pos in sorted(inline.items()):
        # A glossary entry licenses earlier use: that is what a glossary is for.
        if abbr in glossary:
            continue
        if first_use.get(abbr, pos) < pos - 2:
            out.append({"level": "warn", "check": "abbreviations",
                        "msg": f"'{abbr}' a definíciója ELŐTT jelenik meg, és nincs "
                               "rövidítésjegyzék sem — az első előfordulásnál kell feloldani"})
    undefined = [a for a, n in used.items()
                 if n >= 3 and a not in inline and a not in glossary
                 and not re.fullmatch(r"[A-Z]{2,3}", a) is None]
    for abbr in sorted(undefined)[:15]:
        out.append({"level": "info", "check": "abbreviations",
                    "msg": f"'{abbr}' {used[abbr]}× használva, de sehol nincs feloldva — "
                           "ellenőrizd, hogy a célújság elvárja-e a feloldást"})
    return out


def check_sentences(sections: dict[str, str]) -> list[dict]:
    out = []
    for name, body in sections.items():
        if name in ("references", "front") or not body:
            continue
        lens = [len(words(s)) for s in sentences(body)]
        if len(lens) < 3:
            continue
        mean = sum(lens) / len(lens)
        long_ = [n for n in lens if n > 45]
        out.append({"level": "info", "check": "sentences",
                    "msg": f"{name}: {len(lens)} mondat, átlag {mean:.0f} szó, "
                           f"45 szó felett {len(long_)}"})
        for s in sentences(body):
            n = len(words(s))
            if n > 60:
                out.append({"level": "warn", "check": "sentences",
                            "msg": f"{name}: {n} szavas mondat — nézd meg, egy vagy két "
                                   f"állítást hordoz-e: \"{s[:90]}…\""})
    return out


def check_passive(sections: dict[str, str]) -> list[dict]:
    out = []
    for name in ("abstract", "introduction", "methods", "results", "discussion",
                 "conclusions"):
        body = sections.get(name)
        if not body:
            continue
        n_words = len(words(body))
        if n_words < 80:
            continue
        rate = 1000 * len(_PASSIVE.findall(body)) / n_words
        note = ""
        if name in ("methods", "results") and rate < 8:
            note = "  (szokatlanul alacsony egy Methods/Results részhez — nem hiba)"
        elif name in ("introduction", "discussion") and rate > 25:
            note = "  (magas egy érvelő részhez — nézd meg, hol veszik el a cselekvő)"
        out.append({"level": "info", "check": "passive",
                    "msg": f"{name}: passzív ~{rate:.0f}/1000 szó{note}"})
    return out


def check_consistency(text: str) -> list[dict]:
    """The three the reference edit deliberately did NOT globalise. Report the split."""
    out = []

    p_low = len(re.findall(r"(?<![A-Za-z])p\s*[=<>≤≥]", text))
    p_up = len(re.findall(r"(?<![A-Za-z])P\s*[=<>≤≥]", text))
    if p_low and p_up:
        out.append({"level": "warn", "check": "consistency",
                    "msg": f"`p` és `P` is szerepel statisztikai értéknél "
                           f"({p_low}× kisbetűs, {p_up}× nagybetűs). Válaszd a célújság "
                           "konvencióját, vidd végig — és MONDD MEG a szerzőnek, hogy "
                           "globálisan írtad át: a tracked change-ben egy néma eset-csere "
                           "úgy néz ki, mint egy eredmény átírása."})

    # The value before ± is very often percent-suffixed in this literature
    # ("12.0% ± 4.3%"), so a bare \d on the left misses exactly the cases that
    # make a manuscript inconsistent.
    pm_spaced = len(re.findall(r"[\d%]\s+±\s+[\d\-+.]", text))
    pm_closed = len(re.findall(r"[\d%]±[\d\-+.]", text))
    if pm_spaced and pm_closed:
        out.append({"level": "warn", "check": "consistency",
                    "msg": f"a `±` {pm_spaced}× szóközzel, {pm_closed}× szóköz nélkül. "
                           "Egy szakaszon belül tedd egységessé; a globális döntést kérdezd meg."})

    op_spaced = len(re.findall(r"\(\s*[<>≤≥]\s+\d", text))
    op_closed = len(re.findall(r"\(\s*[<>≤≥]\d", text))
    if op_spaced and op_closed:
        out.append({"level": "warn", "check": "consistency",
                    "msg": f"zárójelben az összehasonlító jel {op_spaced}× szóközzel, "
                           f"{op_closed}× anélkül — a referencia-szerkesztés zárójelben "
                           "összezárta: `(>18%)`"})

    for m in re.finditer(r"\b(\d+(?:\.\d+)?)(mg|kg|ml|mL|g|µg|ug|mmHg|mm|cm|s|min|h)\b", text):
        out.append({"level": "warn", "check": "units",
                    "msg": f"hiányzó szóköz érték és mértékegység között: "
                           f"`{m.group(0)}` → `{m.group(1)} {m.group(2)}`"})

    for m in re.finditer(r"\b(\d+(?:st|nd|rd|th))([a-z]+)", text):
        # Show the whole run-together token: `14thd` sends the reader hunting,
        # `14thday` is findable with one search.
        out.append({"level": "warn", "check": "units",
                    "msg": f"összeragadt sorszám: `{m.group(0)}` → "
                           f"`{m.group(1)} {m.group(2)}`"})

    lower_ref = re.findall(r"(?<![A-Za-z])(table|figure|fig\.)\s+\d", text)
    if lower_ref:
        out.append({"level": "warn", "check": "references-to-floats",
                    "msg": f"{len(lower_ref)}× kisbetűs `table`/`figure` + szám — "
                           "számmal hivatkozva nagybetűs: `Table 2`, `Figure 2 a-c`"})
    return out


def check_spelling(text: str) -> list[dict]:
    low = text.lower()
    uk = [(u, s) for u, s in UK_US if re.search(rf"\b{u}", low)]
    us = [(u, s) for u, s in UK_US if re.search(rf"\b{s}", low)]
    out = []
    if uk and us:
        out.append({"level": "warn", "check": "spelling",
                    "msg": f"UK és US helyesírás keverve — UK: "
                           f"{', '.join(u for u, _ in uk[:6])} | US: "
                           f"{', '.join(s for _, s in us[:6])}. Egy kéziratban egy konvenció."})
    elif uk:
        out.append({"level": "info", "check": "spelling",
                    "msg": f"UK helyesírás ({', '.join(u for u, _ in uk[:6])}) — "
                           "ellenőrizd, hogy a célújság ezt kéri-e"})
    return out


def check_register(text: str) -> list[dict]:
    out = []
    for pattern, note in REGISTER_MARKERS:
        for m in re.finditer(pattern, text, re.I):
            start = max(0, m.start() - 40)
            out.append({"level": "warn", "check": "register",
                        "msg": f"`{m.group(0)}` — {note}\n      …"
                               f"{text[start:m.end() + 40].strip()}…"})
    return out


CHECKS = {
    "counts": lambda t, s: check_counts(s),
    "abbreviations": lambda t, s: check_abbreviations(t),
    "sentences": lambda t, s: check_sentences(s),
    "passive": lambda t, s: check_passive(s),
    "consistency": lambda t, s: check_consistency(t),
    "spelling": lambda t, s: check_spelling(t),
    "register": lambda t, s: check_register(t),
}


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("manuscript", type=Path)
    ap.add_argument("--json", type=Path, help="machine-readable report")
    ap.add_argument("--only", help="comma-separated subset: " + ", ".join(CHECKS))
    args = ap.parse_args(argv)

    if not args.manuscript.exists():
        sys.exit(f"nincs ilyen fájl: {args.manuscript}")

    text = read_text(args.manuscript)
    sections = split_sections(text)

    wanted = ([c.strip() for c in args.only.split(",")] if args.only else list(CHECKS))
    unknown = [c for c in wanted if c not in CHECKS]
    if unknown:
        sys.exit(f"ismeretlen ellenőrzés: {', '.join(unknown)}\n  van: {', '.join(CHECKS)}")

    findings: list[dict] = []
    for name in wanted:
        findings.extend(CHECKS[name](text, sections))

    print(f"MANUSCRIPT CHECK — {args.manuscript.name}")
    print(f"felismert szakaszok: {', '.join(k for k in sections if k != 'front') or 'nincs'}")
    print("=" * 70)
    for level in ("warn", "info"):
        rows = [f for f in findings if f["level"] == level]
        if not rows:
            continue
        print(f"\n{'JAVÍTANDÓ' if level == 'warn' else 'MÉRÉS'} ({len(rows)})")
        for f in rows:
            print(f"  [{f['check']}] {f['msg']}")

    n_warn = sum(1 for f in findings if f["level"] == "warn")
    print(f"\n{n_warn} javítandó tétel. Ezek mechanikusak — javítsd őket, mielőtt "
          "stílusra olvasnál.")

    if args.json:
        args.json.write_text(json.dumps(
            {"file": str(args.manuscript), "sections": list(sections),
             "findings": findings}, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"JSON: {args.json}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
