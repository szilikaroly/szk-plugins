"""Extract text and coarse structure from a manuscript file.

Returns a Document with:
  .paragraphs : list of (text, is_heading_style)   -> feeds section splitting
  .text       : the whole plain text
  .sections   : dict from pc_lib.split_sections
  .source     : file type used

Prefers python-docx for .docx (keeps heading styles), PyMuPDF/pdftotext for PDF,
latextotext for .tex, and reads .txt/.md directly.
"""
from __future__ import annotations

import subprocess
from pathlib import Path

import pc_lib


class Document:
    def __init__(self, paragraphs, source):
        self.paragraphs = paragraphs
        self.source = source
        self.text = "\n".join(t for t, _ in paragraphs)
        self.sections = pc_lib.split_sections(paragraphs)


def _from_docx(path):
    import docx
    d = docx.Document(str(path))
    paras = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        is_head = style.startswith("heading") or style == "title"
        paras.append((t, is_head))
    return paras


def _from_pdf(path):
    text = ""
    try:
        import fitz
        doc = fitz.open(str(path))
        text = "\n".join(pg.get_text() for pg in doc)
    except Exception:
        try:
            text = subprocess.run(["pdftotext", "-layout", str(path), "-"],
                                  capture_output=True, text=True).stdout
        except Exception:
            text = ""
    return _plain_to_paras(text)


def _from_tex(path):
    try:
        out = subprocess.run(["latextotext", str(path)],
                             capture_output=True, text=True).stdout
        if out.strip():
            return _plain_to_paras(out)
    except Exception:
        pass
    return _plain_to_paras(Path(path).read_text(errors="ignore"))


def _plain_to_paras(text):
    paras = []
    for line in (text or "").splitlines():
        t = line.strip()
        if t:
            paras.append((t, False))
    return paras


def load(path) -> Document:
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"manuscript not found: {p}")
    suf = p.suffix.lower()
    if suf == ".docx":
        return Document(_from_docx(p), "docx")
    if suf == ".pdf":
        return Document(_from_pdf(p), "pdf")
    if suf in (".tex", ".latex"):
        return Document(_from_tex(p), "tex")
    if suf in (".txt", ".md", ".text"):
        return Document(_plain_to_paras(p.read_text(errors="ignore")), suf[1:])
    # try doctotext as a last resort (handles .doc, .odt, .rtf ...)
    try:
        out = subprocess.run(["doctotext", str(p)],
                             capture_output=True, text=True).stdout
        if out.strip():
            return Document(_plain_to_paras(out), "doctotext")
    except Exception:
        pass
    raise ValueError(f"unsupported file type: {suf}")
